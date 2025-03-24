from datetime import datetime, timedelta
import os
import random
import logging
import json
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.utils import formatdate
from docx import Document
from docx.shared import Pt
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

# Load configuration from config.json
try:
    with open("config_advanced.json", "r") as config_file:
        config = json.load(config_file)
    
    # Extract configuration sections
    request_types = config["request_types"]
    subjects = config["subjects"]
    bodies = config["bodies"]
    bodies_no_request = config["bodies_no_request"]
    bank_names = config["bank_names"]
    random_names = config["random_names"]
    date_range = config["date_range"]
    amount_range = config["amount_range"]
    attachment_details = config["attachment_details"]
    paths = config["paths"]
    email_chain_config = config["email_chain"]
    processing = config["processing"]
    email_types = config["email_types"]  # Load email_types from config

    # Extract specific configuration values
    email_files_dir = paths["email_files_dir"]
    email_generation_log_file = paths["email_generation_log_file"]
    num_emails_per_type = processing["num_emails_per_type"]
    num_layers = processing["num_layers"]
    num_duplicate_emails = processing["num_duplicate_emails"]
    deal_id_min = processing["deal_id_min"]
    deal_id_max = processing["deal_id_max"]
    amount_min = processing["amount_min"]
    amount_max = processing["amount_max"]
    sender_email_prefix = email_chain_config["sender_email_prefix"]
    sender_email_domain = email_chain_config["sender_email_domain"]
    recipient_email = email_chain_config["recipient_email"]
    reply_prefix = email_chain_config["reply_prefix"]
    reply_body = email_chain_config["reply_body"]

except FileNotFoundError:
    logging.error("Configuration file 'config.json' not found.")
    raise FileNotFoundError("Configuration file 'config.json' not found.")
except json.JSONDecodeError as e:
    logging.error(f"Error decoding config.json: {e}")
    raise json.JSONDecodeError(f"Error decoding config.json: {e}")
except KeyError as e:
    logging.error(f"Missing required key in config.json: {e}")
    raise KeyError(f"Missing required key in config.json: {e}")

# Set up logging for email generation
logging.basicConfig(
    filename=email_generation_log_file,
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# Function to generate a random name (for borrower and account name)
def generate_random_name():
    try:
        first = random_names["first"]
        last = random_names["last"]
        return f"{random.choice(first)}{random.choice(last)}"
    except Exception as e:
        logging.error(f"Error generating random name: {e}")
        return "UnknownCorp"

# Function to generate a random date between start_date and end_date
def generate_random_date():
    try:
        start_date = datetime.strptime(date_range["start_date"], "%Y-%m-%d")
        end_date = datetime.strptime(date_range["end_date"], "%Y-%m-%d")
        delta = end_date - start_date
        random_days = random.randint(0, delta.days)
        random_date = start_date + timedelta(days=random_days)
        return random_date.strftime("%d-%b-%Y")
    except Exception as e:
        logging.error(f"Error generating random date: {e}")
        return "01-Jan-2025"

# Function to generate a random amount
def generate_random_amount():
    try:
        min_amount = amount_range["min_amount"]
        max_amount = amount_range["max_amount"]
        amount = random.uniform(min_amount, max_amount)
        return round(amount, 3)
    except Exception as e:
        logging.error(f"Error generating random amount: {e}")
        return 1000000.0

# Function to create a PDF with financial content
def create_pdf_attachment(deal_id, request_type, sub_request_type):
    try:
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter

        pdf_config = attachment_details["pdf"]
        title_position = pdf_config["title_position"]
        content_position = pdf_config["content_position"]
        footer_position = pdf_config["footer_position"]

        # Add a title
        c.setFont(pdf_config["title_font"], pdf_config["title_font_size"])
        c.drawString(title_position[0] * inch, height + title_position[1] * inch, pdf_config["title"])

        # Add financial content
        c.setFont(pdf_config["content_font"], pdf_config["content_font_size"])
        text = c.beginText(content_position[0] * inch, height + content_position[1] * inch)
        
        borrower = generate_random_name()
        deal_name = f"Deal-{deal_id}"
        effective_date = generate_random_date()
        original_amount = generate_random_amount()
        adjusted_amount = original_amount + random.uniform(
            amount_range["adjustment_increment_min"],
            amount_range["adjustment_increment_max"]
        )
        bank_name = random.choice(bank_names)
        account_number = random.randint(100000000, 999999999)
        account_name = generate_random_name()

        financial_details = [
            f"Request Type: {request_type}",
            f"Description: {sub_request_type}",
            f"Borrower: {borrower}",
            f"Deal Name: {deal_name}",
            f"Effective {effective_date}, the lender shares of facility term loan A2 have been adjusted",
            f"your share of commitment was USD {original_amount:,.3f}, it has been increased to USD {adjusted_amount:,.3f}",
            "",
            f"Bank Name: {bank_name}",
            f"Account #: {account_number}",
            f"Account Name: {account_name}"
        ]

        for line in financial_details:
            text.textLine(line)
        c.drawText(text)

        # Add a footer
        c.setFont(pdf_config["footer_font"], pdf_config["footer_font_size"])
        c.drawString(footer_position[0] * inch, footer_position[1] * inch, pdf_config["footer"])

        c.showPage()
        c.save()
        buffer.seek(0)
        return buffer, f"{pdf_config['filename_prefix']}{deal_id}{pdf_config['filename_suffix']}"
    except Exception as e:
        logging.error(f"Error creating PDF attachment for deal_{deal_id}: {e}")
        return None, None

# Function to create a DOCX with financial content
def create_doc_attachment(deal_id, request_type, sub_request_type):
    try:
        doc = Document()
        docx_config = attachment_details["docx"]

        # Add a title
        doc.add_heading(docx_config["title"], level=docx_config["title_level"])
        
        # Add financial content
        borrower = generate_random_name()
        deal_name = f"Deal-{deal_id}"
        effective_date = generate_random_date()
        original_amount = generate_random_amount()
        adjusted_amount = original_amount + random.uniform(
            amount_range["adjustment_increment_min"],
            amount_range["adjustment_increment_max"]
        )
        bank_name = random.choice(bank_names)
        account_number = random.randint(100000000, 999999999)
        account_name = generate_random_name()

        financial_details = [
            f"Request Type: {request_type}",
            f"Description: {sub_request_type}",
            f"Borrower: {borrower}",
            f"Deal Name: {deal_name}",
            f"Effective {effective_date}, the lender shares of facility term loan A2 have been adjusted",
            f"your share of commitment was USD {original_amount:,.3f}, it has been increased to USD {adjusted_amount:,.3f}",
            "",
            f"Bank Name: {bank_name}",
            f"Account #: {account_number}",
            f"Account Name: {account_name}"
        ]

        for detail in financial_details:
            doc.add_paragraph(detail)
        
        # Add a footer as an italicized paragraph
        footer = doc.add_paragraph(docx_config["footer"])
        footer.runs[0].font.italic = docx_config["footer_italic"]
        footer.runs[0].font.size = Pt(docx_config["footer_font_size"])

        # Save the document to a buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer, f"{docx_config['filename_prefix']}{deal_id}{docx_config['filename_suffix']}"
    except Exception as e:
        logging.error(f"Error creating DOCX attachment for deal_{deal_id}: {e}")
        return None, None

# Function to create an email chain
def create_email_chain(email_id, num_layers, attachment_layer=None, attachment_type=None, include_request_in_body=True, custom_body=None):
    try:
        msg = MIMEMultipart()
        msg['From'] = f"{sender_email_prefix}{email_id}{sender_email_domain}"
        msg['To'] = recipient_email
        msg['Date'] = formatdate(localtime=True)

        # Randomly select request type and sub-request type
        request_type = random.choice(list(request_types.keys()))
        sub_request_type = random.choice(request_types[request_type])
        deal_id = random.randint(deal_id_min, deal_id_max)
        amount = random.randint(amount_min, amount_max)

        # Set subject and body
        subject = random.choice(subjects).format(request_type=request_type, sub_request_type=sub_request_type)
        if include_request_in_body:
            body = random.choice(bodies).format(
                request_type=request_type,
                sub_request_type=sub_request_type,
                deal_id=deal_id,
                amount=amount,
                expiration_date=date_range["expiration_date"]
            )
        else:
            body = random.choice(bodies_no_request).format(
                deal_id=deal_id,
                amount=amount,
                expiration_date=date_range["expiration_date"]
            )
        if custom_body:
            body = custom_body

        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain'))

        # Add attachments on the specified layer
        current_layer = 1
        current_msg = msg
        current_body = body

        for layer in range(1, num_layers + 1):
            if layer == attachment_layer and attachment_type:
                if "pdf" in attachment_type:
                    buffer, filename = create_pdf_attachment(deal_id, request_type, sub_request_type)
                    if buffer and filename:
                        attachment = MIMEApplication(buffer.read(), _subtype="pdf")
                        attachment.add_header('Content-Disposition', 'attachment', filename=filename)
                        current_msg.attach(attachment)
                if "doc" in attachment_type:
                    buffer, filename = create_doc_attachment(deal_id, request_type, sub_request_type)
                    if buffer and filename:
                        attachment = MIMEApplication(buffer.read(), _subtype="doc")
                        attachment.add_header('Content-Disposition', 'attachment', filename=filename)
                        current_msg.attach(attachment)

            # Create the next layer in the email chain
            if layer < num_layers:
                reply_msg = MIMEMultipart()
                reply_msg['From'] = recipient_email
                reply_msg['To'] = f"{sender_email_prefix}{email_id}{sender_email_domain}"
                reply_msg['Date'] = formatdate(localtime=True)
                reply_msg['Subject'] = f"{reply_prefix}{subject}"
                reply_msg.attach(MIMEText(reply_body.format(request_type=request_type, body=current_body), 'plain'))
                current_body = reply_msg.as_string()
                current_msg = reply_msg

        return current_msg.as_string(), request_type, sub_request_type
    except Exception as e:
        logging.error(f"Error creating email chain for email_{email_id}: {e}")
        return None, None, None

# Function to generate emails for a specific type
def generate_emails(type_id, num_emails, num_layers, attachment_layer=None, attachment_type=None, include_request_in_body=True, custom_emails=None, randomize=False):
    type_dir = f"{email_files_dir}/type_{type_id}"
    os.makedirs(type_dir, exist_ok=True)
    generated_emails = []

    for i in range(num_emails):
        try:
            email_id = f"{type_id}_{i}"
            # Randomize attachment layer and type if specified
            current_attachment_layer = attachment_layer
            current_attachment_type = attachment_type
            if randomize:
                current_attachment_layer = random.randint(1, num_layers)
                current_attachment_type = random.choice([["pdf"], ["doc"], ["pdf", "doc"]])

            if custom_emails and i < len(custom_emails):
                eml_content, request_type, sub_request_type = custom_emails[i]
            else:
                eml_content, request_type, sub_request_type = create_email_chain(
                    email_id,
                    num_layers,
                    attachment_layer=current_attachment_layer,
                    attachment_type=current_attachment_type,
                    include_request_in_body=include_request_in_body
                )

            if eml_content:
                file_path = f"{type_dir}/email_{email_id}.eml"
                with open(file_path, "w") as f:
                    f.write(eml_content)
                logging.info(f"Generated email_{email_id}.eml in type_{type_id} with Request Type: {request_type}, Sub-Request Type: {sub_request_type}")
                generated_emails.append((eml_content, request_type, sub_request_type))
        except Exception as e:
            logging.error(f"Error generating email_{email_id} for type_{type_id}: {e}")
            continue

    return generated_emails

# Generate emails for all specified types
total_emails = 0
for email_type in email_types[:-1]:  # Exclude type 17 for now
    type_id = email_type["type_id"]
    generated_emails = generate_emails(
        type_id=type_id,
        num_emails=num_emails_per_type,
        num_layers=num_layers,
        attachment_layer=email_type["attachment_layer"],
        attachment_type=email_type["attachment_type"],
        include_request_in_body=email_type["include_request_in_body"]
    )
    total_emails += len(generated_emails)

# Generate emails for type 17 (50 random + 50 duplicates)
type_17_config = email_types[-1]
type_17_num_emails = type_17_config.get("num_emails", num_duplicate_emails)  # Use configured num_emails or default
type_17_emails = generate_emails(
    type_id=type_17_config["type_id"],
    num_emails=type_17_num_emails,
    num_layers=num_layers,
    attachment_layer=type_17_config["attachment_layer"],
    attachment_type=type_17_config["attachment_type"],
    include_request_in_body=type_17_config["include_request_in_body"],
    randomize=type_17_config.get("randomize", False)
)

# Generate duplicates for type 17
type_17_duplicates = generate_emails(
    type_id=type_17_config["type_id"],
    num_emails=type_17_num_emails,
    num_layers=num_layers,
    attachment_layer=type_17_config["attachment_layer"],
    attachment_type=type_17_config["attachment_type"],
    include_request_in_body=type_17_config["include_request_in_body"],
    custom_emails=type_17_emails,  # Use the same emails to create duplicates
    randomize=type_17_config.get("randomize", False)
)

total_emails += len(type_17_emails) + len(type_17_duplicates)

print(f"Generated {total_emails} .eml files in '{email_files_dir}' directory. Check {email_generation_log_file} for details.")