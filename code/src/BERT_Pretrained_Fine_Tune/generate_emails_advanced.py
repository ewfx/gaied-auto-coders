from datetime import datetime, timedelta
import os
import random
import logging
import json
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.utils import formatdate
from docx import Document
from docx.shared import Pt
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

# Load configuration from config.json
try:
    with open("config_advanced.json", "r") as config_file:
        config = json.load(config_file)
    
    # Extract configuration sections
    request_types = config["request_types"]
    subjects = config["subjects"]
    bodies_no_request = config["bodies_no_request"]
    bank_names = config["bank_names"]
    random_names = config["random_names"]
    date_range = config["date_range"]
    amount_range = config["amount_range"]
    attachment_details = config["attachment_details"]
    paths = config["paths"]
    email_chain_config = config["email_chain"]
    processing = config["processing"]
    email_types = config["email_types"]
    email_body_templates = config["email_body_templates"]
    attachment_fields = config["attachment_fields"]
    expiration_date_range = config["expiration_date_range"]

    # Extract specific configuration values
    email_files_dir = paths["email_files_dir"]
    email_generation_log_file = paths["email_generation_log_file"]
    num_emails_per_type = processing["num_emails_per_type"]
    num_layers = processing["num_layers"]
    num_duplicate_emails = processing["num_duplicate_emails"]
    deal_id_min = processing["deal_id_min"]
    deal_id_max = processing["deal_id_max"]
    sender_email_prefix = email_chain_config["sender_email_prefix"]
    sender_email_domain = email_chain_config["sender_email_domain"]
    recipient_email = email_chain_config["recipient_email"]
    reply_prefix = email_chain_config["reply_prefix"]
    reply_body = email_chain_config["reply_body"]
except FileNotFoundError:
    logging.error("Configuration file 'config_advanced.json' not found.")
    raise FileNotFoundError("Configuration file 'config_advanced.json' not found.")
except json.JSONDecodeError as e:
    logging.error(f"Error decoding config_advanced.json: {e}")
    raise json.JSONDecodeError(f"Error decoding config_advanced.json: {e}")
except KeyError as e:
    logging.error(f"Missing required key in config_advanced.json: {e}")
    raise KeyError(f"Missing required key in config_advanced.json: {e}")

# Set up logging for email generation
logging.basicConfig(
    filename=email_generation_log_file,
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# Utility Functions
def generate_random_name():
    try:
        first = random_names["first"]
        last = random_names["last"]
        return f"{random.choice(first)}{random.choice(last)}"
    except Exception as e:
        logging.error(f"Error generating random name: {e}")
        return "UnknownCorp"

def generate_random_date():
    try:
        start_date = datetime.strptime(date_range["start_date"], "%Y-%m-%d")
        end_date = datetime.strptime(date_range["end_date"], "%Y-%m-%d")
        delta = end_date - start_date
        random_days = random.randint(0, delta.days)
        random_date = start_date + timedelta(days=random_days)
        return random_date.strftime("%d-%b-%Y")
    except Exception as e:
        logging.error(f"Error generating random date: {e}")
        return "01-Jan-2025"

def generate_random_amount():
    try:
        min_amount = amount_range["min_amount"]
        max_amount = amount_range["max_amount"]
        amount = random.uniform(min_amount, max_amount)
        return round(amount, 3)
    except Exception as e:
        logging.error(f"Error generating random amount: {e}")
        return 1000000.0

# Attachment Creation Functions
def create_pdf_attachment(deal_id, requests, include_fields=True):
    try:
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter

        pdf_config = attachment_details["pdf"]
        title_position = pdf_config["title_position"]
        content_position = pdf_config["content_position"]
        footer_position = pdf_config["footer_position"]

        c.setFont(pdf_config["title_font"], pdf_config["title_font_size"])
        c.drawString(title_position[0] * inch, height + title_position[1] * inch, pdf_config["title"])

        c.setFont(pdf_config["content_font"], pdf_config["content_font_size"])
        text = c.beginText(content_position[0] * inch, height + content_position[1] * inch)
        
        if include_fields:
            borrower = generate_random_name()
            deal_name = f"Deal-{deal_id}"
            effective_date = generate_random_date()
            expiration_days = random.randint(
                expiration_date_range["min_days_after_effective"],
                expiration_date_range["max_days_after_effective"]
            )
            expiration_date = (datetime.strptime(effective_date, "%d-%b-%Y") + 
                              timedelta(days=expiration_days)).strftime("%d-%b-%Y")
            original_amount = generate_random_amount()
            adjusted_amount = original_amount + random.uniform(
                amount_range["adjustment_increment_min"],
                amount_range["adjustment_increment_max"]
            )
            bank_name = random.choice(bank_names)
            account_number = random.randint(100000000, 999999999)
            account_name = generate_random_name()

            attachment_fields_pdf = attachment_fields["pdf"]
            financial_details = []
            for field in attachment_fields_pdf:
                if field == "Request Type":
                    for req in requests:
                        financial_details.append(f"Request Type: {req['request_type']}")
                        financial_details.append(f"Description: {req['sub_request_type']}")
                elif field == "Borrower":
                    financial_details.append(f"Borrower: {borrower}")
                elif field == "Deal Name":
                    financial_details.append(f"Deal Name: {deal_name}")
                elif field == "Effective Date":
                    financial_details.append(f"Effective Date: {effective_date}")
                elif field == "Expiration Date":
                    financial_details.append(f"Expiration Date: {expiration_date}")
                elif field == "Original Amount":
                    financial_details.append(f"Original Amount: $ {original_amount:,.3f}")
                elif field == "Adjusted Amount":
                    financial_details.append(f"Adjusted Amount: $ {adjusted_amount:,.3f}")
                elif field == "Bank Name":
                    financial_details.append(f"Bank Name: {bank_name}")
                elif field == "Account #":
                    financial_details.append(f"Account #: {account_number}")
                elif field == "Account Name":
                    financial_details.append(f"Account Name: {account_name}")
        else:
            document_type = "Metadata Document"
            generated_by = generate_random_name()
            generation_date = datetime.now().strftime("%d-%b-%Y")
            document_id = f"DOC-{random.randint(1000, 9999)}"
            financial_details = [
                f"Document Type: {document_type}",
                f"Generated By: {generated_by}",
                f"Generation Date: {generation_date}",
                f"Document ID: {document_id}",
                "This document contains metadata information only.",
                "For detailed financial information, refer to the accompanying attachment."
            ]

        for line in financial_details:
            text.textLine(line)
        c.drawText(text)

        c.setFont(pdf_config["footer_font"], pdf_config["footer_font_size"])
        c.drawString(footer_position[0] * inch, footer_position[1] * inch, pdf_config["footer"])

        c.showPage()
        c.save()
        buffer.seek(0)
        return buffer, f"{pdf_config['filename_prefix']}{deal_id}{pdf_config['filename_suffix']}"
    except Exception as e:
        logging.error(f"Error creating PDF attachment for deal_{deal_id}: {e}")
        return None, None

def create_doc_attachment(deal_id, requests, include_fields=True):
    try:
        doc = Document()
        docx_config = attachment_details["docx"]

        doc.add_heading(docx_config["title"], level=docx_config["title_level"])
        
        if include_fields:
            borrower = generate_random_name()
            deal_name = f"Deal-{deal_id}"
            effective_date = generate_random_date()
            expiration_days = random.randint(
                expiration_date_range["min_days_after_effective"],
                expiration_date_range["max_days_after_effective"]
            )
            expiration_date = (datetime.strptime(effective_date, "%d-%b-%Y") + 
                              timedelta(days=expiration_days)).strftime("%d-%b-%Y")
            original_amount = generate_random_amount()
            adjusted_amount = original_amount + random.uniform(
                amount_range["adjustment_increment_min"],
                amount_range["adjustment_increment_max"]
            )
            bank_name = random.choice(bank_names)
            account_number = random.randint(100000000, 999999999)
            account_name = generate_random_name()

            attachment_fields_docx = attachment_fields["docx"]
            financial_details = []
            for field in attachment_fields_docx:
                if field == "Request Type":
                    for req in requests:
                        financial_details.append(f"Request Type: {req['request_type']}")
                        financial_details.append(f"Description: {req['sub_request_type']}")
                elif field == "Borrower":
                    financial_details.append(f"Borrower: {borrower}")
                elif field == "Deal Name":
                    financial_details.append(f"Deal Name: {deal_name}")
                elif field == "Effective Date":
                    financial_details.append(f"Effective Date: {effective_date}")
                elif field == "Expiration Date":
                    financial_details.append(f"Expiration Date: {expiration_date}")
                elif field == "Original Amount":
                    financial_details.append(f"Original Amount: $ {original_amount:,.3f}")
                elif field == "Adjusted Amount":
                    financial_details.append(f"Adjusted Amount: $ {adjusted_amount:,.3f}")
                elif field == "Bank Name":
                    financial_details.append(f"Bank Name: {bank_name}")
                elif field == "Account #":
                    financial_details.append(f"Account #: {account_number}")
                elif field == "Account Name":
                    financial_details.append(f"Account Name: {account_name}")
        else:
            document_type = "Metadata Document"
            generated_by = generate_random_name()
            generation_date = datetime.now().strftime("%d-%b-%Y")
            document_id = f"DOC-{random.randint(1000, 9999)}"
            financial_details = [
                f"Document Type: {document_type}",
                f"Generated By: {generated_by}",
                f"Generation Date: {generation_date}",
                f"Document ID: {document_id}",
                "This document contains metadata information only.",
                "For detailed financial information, refer to the accompanying attachment."
            ]

        for detail in financial_details:
            doc.add_paragraph(detail)
        
        footer = doc.add_paragraph(docx_config["footer"])
        footer.runs[0].font.italic = docx_config["footer_italic"]
        footer.runs[0].font.size = Pt(docx_config["footer_font_size"])

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer, f"{docx_config['filename_prefix']}{deal_id}{docx_config['filename_suffix']}"
    except Exception as e:
        logging.error(f"Error creating DOCX attachment for deal_{deal_id}: {e}")
        return None, None

# Modified Email Chain Creation Function
def create_email_chain(email_id, num_layers, attachment_layer=None, attachment_type=None, 
                      include_request_in_body=True, custom_body=None):
    try:
        msg = MIMEMultipart()
        msg['From'] = f"{sender_email_prefix}{email_id}{sender_email_domain}"
        msg['To'] = recipient_email
        msg['Date'] = formatdate(localtime=True)

        # Determine number of requests based on attachment_layer (default to 1 if None)
        num_requests = min(max(1, attachment_layer if attachment_layer is not None else 1), 5)  # 1 to 5 requests
        selected_requests = random.sample(list(request_types.keys()), num_requests)
        requests = []
        for req in selected_requests:
            sub_req = random.choice(request_types[req])
            requests.append({"request_type": req, "sub_request_type": sub_req})
        primary_request = random.choice(selected_requests)
        primary_sub_request = next(req['sub_request_type'] for req in requests if req['request_type'] == primary_request)
        deal_id = random.randint(deal_id_min, deal_id_max)

        subject = random.choice(subjects).format(request_type=primary_request, sub_request_type=primary_sub_request)
        
        # Initial email body
        if include_request_in_body:
            body = "I need to discuss the following requests:\n"
            for req in requests:
                body += f"- {req['request_type']} for {req['sub_request_type']}\n"
            body += f"My main concern is {primary_request}.\n"
            body += "Please find the detailed information in the attached document." if attachment_type and attachment_layer == 1 else ""
        else:
            body = random.choice(bodies_no_request)
        if custom_body:
            body = custom_body

        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain'))

        current_layer = 1
        current_msg = msg
        current_body = body

        # Build the email chain with nested replies
        for layer in range(1, num_layers + 1):
            # Add attachments if attachment_type is specified and we're at the attachment_layer
            if attachment_type and layer == attachment_layer:
                # Determine which attachment includes financial fields
                if "pdf" in attachment_type and "doc" in attachment_type:
                    include_fields_in = random.choice(["pdf", "doc"])
                    include_fields_pdf = (include_fields_in == "pdf")
                    include_fields_doc = (include_fields_in == "doc")
                    logging.info(f"For email_{email_id}, fields are included in {include_fields_in.upper()}")
                else:
                    include_fields_pdf = "pdf" in attachment_type
                    include_fields_doc = "doc" in attachment_type

                # Attach PDF if specified in attachment_type
                if "pdf" in attachment_type:
                    buffer, filename = create_pdf_attachment(deal_id, requests, include_fields=include_fields_pdf)
                    if buffer and filename:
                        attachment = MIMEApplication(buffer.read(), _subtype="pdf")
                        attachment.add_header('Content-Disposition', 'attachment', filename=filename)
                        current_msg.attach(attachment)
                        logging.info(f"Attached PDF: {filename} at layer {layer} for email_{email_id} (fields: {include_fields_pdf})")

                # Attach DOCX if specified in attachment_type
                if "doc" in attachment_type:
                    buffer, filename = create_doc_attachment(deal_id, requests, include_fields=include_fields_doc)
                    if buffer and filename:
                        attachment = MIMEApplication(buffer.read(), _subtype="docx")
                        attachment.add_header('Content-Disposition', 'attachment', filename=filename)
                        current_msg.attach(attachment)
                        logging.info(f"Attached DOCX: {filename} at layer {layer} for email_{email_id} (fields: {include_fields_doc})")

            # Add reply emails if not at the last layer
            elif layer < num_layers and (layer <= attachment_layer or attachment_layer is None):
                reply_msg = MIMEMultipart()
                reply_msg['From'] = recipient_email
                reply_msg['To'] = f"{sender_email_prefix}{email_id}{sender_email_domain}"
                reply_msg['Date'] = formatdate(localtime=True)
                reply_msg['Subject'] = f"{reply_prefix}{subject}"

                # Construct reply body with nested "Original Message" structure
                reply_text = f"Replying to your email regarding {primary_request}...\n\nOriginal Message:\n{current_body}"
                reply_msg.attach(MIMEText(reply_text, 'plain'))
                
                current_body = reply_msg.as_string()
                current_msg = reply_msg
                current_layer += 1

        return current_msg.as_string(), primary_request, primary_sub_request
    except Exception as e:
        logging.error(f"Error creating email chain for email_{email_id}: {e}")
        return None, None, None

# Modified Email Generation Function
def generate_emails(type_id, num_emails, num_layers, attachment_layer=None, attachment_type=None, 
                   include_request_in_body=True, custom_emails=None, randomize=False):
    type_dir = f"{email_files_dir}/type_{type_id}"
    os.makedirs(type_dir, exist_ok=True)
    generated_emails = []

    for i in range(num_emails):
        try:
            email_id = f"{type_id}_{i}"
            current_attachment_layer = attachment_layer
            current_attachment_type = attachment_type
            
            # Ensure attachment for all emails if attachment_type is not None
            if current_attachment_type is not None:
                if current_attachment_layer is None:
                    current_attachment_layer = 1  # Default to layer 1 if not specified
            # Randomize attachment only if attachment_type is None and randomize is True
            elif randomize:
                current_attachment_layer = random.randint(1, num_layers)
                current_attachment_type = random.choice([["pdf"], ["doc"], ["pdf", "doc"]])
                logging.info(f"Randomized attachment for email_{email_id}: layer={current_attachment_layer}, type={current_attachment_type}")

            if custom_emails and i < len(custom_emails):
                eml_content, request_type, sub_request_type = custom_emails[i]
            else:
                eml_content, request_type, sub_request_type = create_email_chain(
                    email_id,
                    num_layers,
                    attachment_layer=current_attachment_layer,
                    attachment_type=current_attachment_type,
                    include_request_in_body=include_request_in_body
                )

            if eml_content:
                file_path = f"{type_dir}/email_{email_id}.eml"
                with open(file_path, "w") as f:
                    f.write(eml_content)
                logging.info(f"Generated email_{email_id}.eml in type_{type_id} with Request Type: {request_type}, Sub-Request Type: {sub_request_type}, Num Requests: {min(max(1, current_attachment_layer if current_attachment_layer else 1), 3)}")
                generated_emails.append((eml_content, request_type, sub_request_type))
        except Exception as e:
            logging.error(f"Error generating email_{email_id} for type_{type_id}: {e}")
            continue

    return generated_emails

# Generate emails for all specified types
total_emails = 0

for email_type in email_types:
    type_id = email_type["type_id"]
    num_emails = email_type.get("num_emails", num_emails_per_type)
    attachment_layer = email_type["attachment_layer"]
    attachment_type = email_type["attachment_type"]
    include_request_in_body = email_type["include_request_in_body"]
    randomize = email_type.get("randomize", False)

    # Generate initial emails
    generated_emails = generate_emails(
        type_id=type_id,
        num_emails=num_emails,
        num_layers=num_layers,
        attachment_layer=attachment_layer,
        attachment_type=attachment_type,
        include_request_in_body=include_request_in_body,
        randomize=randomize
    )
    total_emails += len(generated_emails)

print(f"Generated {total_emails} .eml files in '{email_files_dir}' directory. Check {email_generation_log_file} for details.")