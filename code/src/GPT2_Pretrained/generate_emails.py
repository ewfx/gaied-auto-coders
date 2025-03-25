from datetime import datetime, timedelta
import os
import random
import logging
import json
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.utils import formatdate
from docx import Document
from docx.shared import Pt
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

# Load configuration from config.json
try:
    with open("config.json", "r") as config_file:
        config = json.load(config_file)
    
    # Extract configuration sections
    request_types = config["request_types"]
    subjects = config["subjects"]
    bodies = config["bodies"]
    bank_names = config["bank_names"]
    random_names = config["random_names"]
    date_range = config["date_range"]
    amount_range = config["amount_range"]
    attachment_details = config["attachment_details"]
    paths = config["paths"]
    email_chain_config = config["email_chain"]
    processing = config["processing"]

    # Extract specific configuration values
    email_files_dir = paths["email_files_dir"]
    email_generation_log_file = paths["email_generation_log_file"]
    num_emails = processing["num_emails"]
    deal_id_min = processing["deal_id_min"]
    deal_id_max = processing["deal_id_max"]
    amount_min = processing["amount_min"]
    amount_max = processing["amount_max"]
    num_emails_in_chain_min = email_chain_config["num_emails_in_chain_min"]
    num_emails_in_chain_max = email_chain_config["num_emails_in_chain_max"]
    sender_email_prefix = email_chain_config["sender_email_prefix"]
    sender_email_domain = email_chain_config["sender_email_domain"]
    recipient_email = email_chain_config["recipient_email"]
    reply_prefix = email_chain_config["reply_prefix"]
    reply_body = email_chain_config["reply_body"]

except FileNotFoundError:
    logging.error("Configuration file 'config.json' not found.")
    raise FileNotFoundError("Configuration file 'config.json' not found.")
except json.JSONDecodeError as e:
    logging.error(f"Error decoding config.json: {e}")
    raise json.JSONDecodeError(f"Error decoding config.json: {e}")
except KeyError as e:
    logging.error(f"Missing required key in config.json: {e}")
    raise KeyError(f"Missing required key in config.json: {e}")

# Set up logging for email generation
logging.basicConfig(
    filename=email_generation_log_file,
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# Function to generate a random name (for borrower and account name)
def generate_random_name():
    first = random_names["first"]
    last = random_names["last"]
    return f"{random.choice(first)}{random.choice(last)}"

# Function to generate a random date between start_date and end_date
def generate_random_date():
    start_date = datetime.strptime(date_range["start_date"], "%Y-%m-%d")
    end_date = datetime.strptime(date_range["end_date"], "%Y-%m-%d")
    delta = end_date - start_date
    random_days = random.randint(0, delta.days)
    random_date = start_date + timedelta(days=random_days)
    return random_date.strftime(date_range["date_format"])  # Use configurable date format

# Function to generate a random amount
def generate_random_amount():
    min_amount = amount_range["min_amount"]
    max_amount = amount_range["max_amount"]
    amount = random.uniform(min_amount, max_amount)
    return round(amount, amount_range["decimal_places"])  # Use configurable decimal places

# Function to create a PDF with financial content
def create_pdf_attachment(deal_id, request_type, sub_request_type, amount, expiration_date):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    pdf_config = attachment_details["pdf"]
    title_position = pdf_config["title_position"]
    content_position = pdf_config["content_position"]
    footer_position = pdf_config["footer_position"]

    # Add a title
    c.setFont(pdf_config["title_font"], pdf_config["title_font_size"])
    c.drawString(title_position[0] * inch, height + title_position[1] * inch, pdf_config["title"])

    # Add financial content
    c.setFont(pdf_config["content_font"], pdf_config["content_font_size"])
    text = c.beginText(content_position[0] * inch, height + content_position[1] * inch)
    
    borrower = generate_random_name()
    deal_name = f"Deal-{deal_id}"
    effective_date = generate_random_date()
    original_amount = generate_random_amount()
    adjusted_amount = original_amount + random.uniform(
        amount_range["adjustment_increment_min"],
        amount_range["adjustment_increment_max"]
    )
    bank_name = random.choice(bank_names)
    account_number = random.randint(100000000, 999999999)
    account_name = generate_random_name()

    financial_details = [
        f"Request Type: {request_type}",
        f"Description: {sub_request_type}",
        f"Borrower: {borrower}",
        f"Deal Name: {deal_name}",
        f"Deal Amount: $ {amount:,.2f}",  # Added Deal Amount
        f"Expiration Date: {expiration_date}",  # Added Expiration Date
        f"Effective {effective_date}, the lender shares of facility term loan A2 have been adjusted",
        f"Your share of commitment was USD {original_amount:,.3f}, it has been increased to USD {adjusted_amount:,.3f}",
        "",
        f"Bank Name: {bank_name}",
        f"Account #: {account_number}",
        f"Account Name: {account_name}"
    ]

    for line in financial_details:
        text.textLine(line)
    c.drawText(text)

    # Add a footer
    c.setFont(pdf_config["footer_font"], pdf_config["footer_font_size"])
    c.drawString(footer_position[0] * inch, footer_position[1] * inch, pdf_config["footer"])

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer, f"{pdf_config['filename_prefix']}{deal_id}{pdf_config['filename_suffix']}"

# Function to create a DOCX with financial content
def create_doc_attachment(deal_id, request_type, sub_request_type, amount, expiration_date):
    doc = Document()
    docx_config = attachment_details["docx"]

    # Add a title
    doc.add_heading(docx_config["title"], level=docx_config["title_level"])
    
    # Add financial content
    borrower = generate_random_name()
    deal_name = f"Deal-{deal_id}"
    effective_date = generate_random_date()
    original_amount = generate_random_amount()
    adjusted_amount = original_amount + random.uniform(
        amount_range["adjustment_increment_min"],
        amount_range["adjustment_increment_max"]
    )
    bank_name = random.choice(bank_names)
    account_number = random.randint(100000000, 999999999)
    account_name = generate_random_name()

    financial_details = [
        f"Request Type: {request_type}",
        f"Description: {sub_request_type}",
        f"Borrower: {borrower}",
        f"Deal Name: {deal_name}",
        f"Deal Amount: $ {amount:,.2f}",  # Added Deal Amount
        f"Expiration Date: {expiration_date}",  # Added Expiration Date
        f"Effective {effective_date}, the lender shares of facility term loan A2 have been adjusted",
        f"Your share of commitment was USD {original_amount:,.3f}, it has been increased to USD {adjusted_amount:,.3f}",
        "",
        f"Bank Name: {bank_name}",
        f"Account #: {account_number}",
        f"Account Name: {account_name}"
    ]

    for detail in financial_details:
        doc.add_paragraph(detail)
    
    # Add a footer as an italicized paragraph
    footer = doc.add_paragraph(docx_config["footer"])
    footer.runs[0].font.italic = docx_config["footer_italic"]
    footer.runs[0].font.size = Pt(docx_config["footer_font_size"])

    # Save the document to a buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, f"{docx_config['filename_prefix']}{deal_id}{docx_config['filename_suffix']}"

# Function to create an email chain
def create_email_chain(email_id, num_emails_in_chain):
    msg = MIMEMultipart()
    msg['From'] = f"{sender_email_prefix}{email_id}{sender_email_domain}"
    msg['To'] = recipient_email
    msg['Date'] = formatdate(localtime=True)

    # Randomly select request type and sub-request type
    request_type = random.choice(list(request_types.keys()))
    sub_request_type = random.choice(request_types[request_type])
    deal_id = random.randint(deal_id_min, deal_id_max)
    amount = random.randint(amount_min, amount_max)
    expiration_date = date_range["expiration_date"]

    # Set subject and body
    subject = random.choice(subjects).format(request_type=request_type, sub_request_type=sub_request_type)
    body = random.choice(bodies).format(
        request_type=request_type,
        sub_request_type=sub_request_type,
        deal_id=deal_id
    )

    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    # Add attachment (PDF or DOC)
    attachment_type = random.choice(["pdf", "doc"])
    if attachment_type == "pdf":
        buffer, filename = create_pdf_attachment(deal_id, request_type, sub_request_type, amount, expiration_date)
    else:
        buffer, filename = create_doc_attachment(deal_id, request_type, sub_request_type, amount, expiration_date)

    attachment = MIMEApplication(buffer.read(), _subtype=attachment_type)
    attachment.add_header('Content-Disposition', 'attachment', filename=filename)
    msg.attach(attachment)

    # Simulate email chain by adding "Re:" or "Fwd:" to subject
    for i in range(num_emails_in_chain - 1):
        reply_msg = MIMEMultipart()
        reply_msg['From'] = recipient_email
        reply_msg['To'] = f"{sender_email_prefix}{email_id}{sender_email_domain}"
        reply_msg['Date'] = formatdate(localtime=True)
        reply_msg['Subject'] = f"{reply_prefix}{subject}"
        reply_msg.attach(MIMEText(reply_body.format(request_type=request_type, body=body), 'plain'))
        body = reply_msg.as_string()

    return msg.as_string(), request_type, sub_request_type

# Generate .eml files
os.makedirs(email_files_dir, exist_ok=True)
for i in range(num_emails):
    num_emails_in_chain = random.randint(num_emails_in_chain_min, num_emails_in_chain_max)
    eml_content, request_type, sub_request_type = create_email_chain(i, num_emails_in_chain)
    file_path = f"{email_files_dir}/email_{i}.eml"
    with open(file_path, "w") as f:
        f.write(eml_content)
    logging.info(f"Generated email_{i}.eml with Request Type: {request_type}, Sub-Request Type: {sub_request_type}")

print(f"Generated {num_emails} .eml files in '{email_files_dir}' directory. Check {email_generation_log_file} for details.")